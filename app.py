import tkinter as tk
from tkinter import filedialog, messagebox
from PIL import Image, ImageTk, ImageDraw
from docx import Document
from docx.shared import Inches
import os

class ImageToPDFConverter:
    def __init__(self, root):
        self.root = root
        self.image_paths = []
        self.output_pdf_name = tk.StringVar()
        self.selected_images_listbox = tk.Listbox(root, selectmode=tk.MULTIPLE, bg="#f0f0f0", fg="#333", font=("Helvetica", 10))

        self.initialize_ui()

    def initialize_ui(self):
        self.root.config(bg="#4a4a4a")

        title_label = tk.Label(self.root, text="Image to PDF Converter", font=("Helvetica", 18, "bold"), bg="#4a4a4a", fg="#ffffff")
        title_label.pack(pady=10)

        creator_label = tk.Label(self.root, text="Creator: Hashan Anjana", font=("Helvetica", 10), bg="#4a4a4a", fg="#cccccc")
        creator_label.pack(pady=5)

        button_frame_top = tk.Frame(self.root, bg="#4a4a4a")
        button_frame_top.pack(pady=10)

        add_image_button = tk.Button(button_frame_top, text="Add Images", command=self.add_images, bg="#66b3ff", fg="#ffffff", font=("Helvetica", 12), relief="flat", padx=10, pady=5)
        add_image_button.pack(side=tk.LEFT, padx=5)

        create_image_button = tk.Button(button_frame_top, text="Create AI Image", command=self.show_text_field, bg="#66b3ff", fg="#ffffff", font=("Helvetica", 12), relief="flat", padx=10, pady=5)
        create_image_button.pack(side=tk.LEFT, padx=5)

        self.selected_images_listbox.pack(pady=5, padx=10, fill=tk.BOTH, expand=True)

        button_frame_bottom = tk.Frame(self.root, bg="#4a4a4a")
        button_frame_bottom.pack(pady=10)

        convert_pdf_button = tk.Button(button_frame_bottom, text="Convert to PDF", command=self.convert_to_pdf, bg="#66b3ff", fg="#ffffff", font=("Helvetica", 12), relief="flat", padx=10, pady=5)
        convert_pdf_button.pack(side=tk.LEFT, padx=5)

        convert_word_button = tk.Button(button_frame_bottom, text="Convert to Word", command=self.convert_to_word, bg="#66b3ff", fg="#ffffff", font=("Helvetica", 12), relief="flat", padx=10, pady=5)
        convert_word_button.pack(side=tk.LEFT, padx=5)

        preview_button = tk.Button(button_frame_bottom, text="Preview Image", command=self.preview_image, bg="#66b3ff", fg="#ffffff", font=("Helvetica", 12), relief="flat", padx=10, pady=5)
        preview_button.pack(side=tk.LEFT, padx=5)

        remove_selected_button = tk.Button(button_frame_bottom, text="Remove Selected", command=self.remove_selected, bg="#ff6666", fg="#ffffff", font=("Helvetica", 12), relief="flat", padx=10, pady=5)
        remove_selected_button.pack(side=tk.LEFT, padx=5)

        clear_all_button = tk.Button(button_frame_bottom, text="Clear All", command=self.clear_all, bg="#ff6666", fg="#ffffff", font=("Helvetica", 12), relief="flat", padx=10, pady=5)
        clear_all_button.pack(side=tk.LEFT, padx=5)

    def add_images(self):
        file_paths = filedialog.askopenfilenames(filetypes=[("Image files", "*.jpg *.jpeg *.png")])
        if file_paths:
            for path in file_paths:
                self.image_paths.append(path)
                self.selected_images_listbox.insert(tk.END, path)

    def convert_to_pdf(self):
        if not self.image_paths:
            messagebox.showerror("Error", "No images selected!")
            return

        images = []
        for path in self.image_paths:
            img = Image.open(path)
            img = img.convert("RGB")
            images.append(img)

        output_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")])
        if output_path:
            images[0].save(output_path, save_all=True, append_images=images[1:])
            messagebox.showinfo("Success", f"PDF saved to {output_path}")

    def convert_to_word(self):
        if not self.image_paths:
            messagebox.showerror("Error", "No images selected!")
            return

        doc = Document()

        for path in self.image_paths:
            doc.add_paragraph(f"Image: {os.path.basename(path)}")
            doc.add_picture(path, width=Inches(6.0))  # Adjust width as needed
            doc.add_page_break()

        output_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if output_path:
            doc.save(output_path)
            messagebox.showinfo("Success", f"Word document saved to {output_path}")

    def preview_image(self):
        selected_indices = self.selected_images_listbox.curselection()
        if not selected_indices:
            messagebox.showerror("Error", "No image selected to preview!")
            return
        selected_index = selected_indices[0]
        selected_image_path = self.image_paths[selected_index]

        preview_window = tk.Toplevel(self.root)
        preview_window.title("Image Preview")
        preview_window.geometry("600x600")

        img = Image.open(selected_image_path)
        img.thumbnail((600, 600))
        img = ImageTk.PhotoImage(img)

        img_label = tk.Label(preview_window, image=img)
        img_label.image = img  # Keep a reference to avoid garbage collection
        img_label.pack(expand=True)

    def remove_selected(self):
        selected_indices = self.selected_images_listbox.curselection()
        if not selected_indices:
            messagebox.showerror("Error", "No images selected to remove!")
            return
        for index in reversed(selected_indices):
            self.selected_images_listbox.delete(index)
            del self.image_paths[index]

    def clear_all(self):
        if not self.image_paths:
            messagebox.showerror("Error", "No images to clear!")
            return
        self.selected_images_listbox.delete(0, tk.END)
        self.image_paths = []

    def show_text_field(self):
        self.ai_image_text = tk.Entry(self.root, width=50, font=("Helvetica", 12))
        self.ai_image_text.pack(pady=10)
        self.ai_image_text.bind("<Return>", self.create_ai_image)

        create_button = tk.Button(self.root, text="Generate Image", command=self.create_ai_image, bg="#66b3ff", fg="#ffffff", font=("Helvetica", 12), relief="flat", padx=10, pady=5)
        create_button.pack(pady=5)

    def create_ai_image(self, event=None):
        text = self.ai_image_text.get()
        if not text:
            messagebox.showerror("Error", "Please enter text to generate AI image!")
            return

        # Dummy implementation for AI image generation
        generated_image_path = f"generated_image_for_{text}.png"

        # Save the generated image to the local disk
        img = Image.new("RGB", (600, 600), color="white")
        d = ImageDraw.Draw(img)
        d.text((10, 10), f"Generated image for: {text}", fill="black")
        img.save(generated_image_path)

        # Add the generated image to the list of images
        self.image_paths.append(generated_image_path)
        self.selected_images_listbox.insert(tk.END, generated_image_path)

        # Display the generated image
        self.display_generated_image(generated_image_path)

    def display_generated_image(self, image_path):
        generated_image_window = tk.Toplevel(self.root)
        generated_image_window.title("Generated AI Image")
        generated_image_window.geometry("600x600")

        img = Image.open(image_path)
        img = img.resize((600, 600))  # Resize to fit the window
        img = ImageTk.PhotoImage(img)

        img_label = tk.Label(generated_image_window, image=img)
        img_label.image = img  # Keep a reference to avoid garbage collection
        img_label.pack(expand=True)

        download_button = tk.Button(generated_image_window, text="Download Image", command=lambda: self.download_image(image_path), bg="#66b3ff", fg="#ffffff", font=("Helvetica", 12), relief="flat", padx=10, pady=5)
        download_button.pack(pady=10)

    def download_image(self, image_path):
        download_path = filedialog.asksaveasfilename(defaultextension=".png", filetypes=[("PNG files", "*.png")])
        if download_path:
            try:
                img = Image.open(image_path)
                img.save(download_path)
                messagebox.showinfo("Success", f"Image downloaded to {download_path}")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to download image: {e}")

def main():
    root = tk.Tk()
    root.title("Image to PDF Converter")
    root.geometry("600x600")
    app = ImageToPDFConverter(root)
    root.mainloop()

if __name__ == "__main__":
    main()
